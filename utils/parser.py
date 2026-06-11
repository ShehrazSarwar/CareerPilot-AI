import os
import io
import pdfplumber
import docx

def extract_text_from_pdf(file_obj) -> str:
    """
    Extracts text from a PDF file object using pdfplumber.
    """
    text_content = []
    try:
        # Streamlit UploadedFile is a file-like object, which pdfplumber can open directly
        with pdfplumber.open(file_obj) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
                else:
                    # In case of scanned pages or image pages, extract_text might return None
                    text_content.append(f"[Page {i+1} has no selectable text]")
        
        full_text = "\n\n".join(text_content).strip()
        return full_text
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")

def extract_text_from_docx(file_obj) -> str:
    """
    Extracts text from a DOCX file object using python-docx.
    """
    try:
        # docx.Document can read from a file-like object (BytesIO)
        doc = docx.Document(file_obj)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text).strip()
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")

def parse_resume(uploaded_file) -> str:
    """
    Routes the uploaded file to the appropriate parser based on its file type/extension.
    """
    if uploaded_file is None:
        raise ValueError("No file provided.")
    
    filename = uploaded_file.name.lower()
    
    # Seek to start of file object to ensure we read from the beginning
    uploaded_file.seek(0)
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(uploaded_file)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        return extract_text_from_docx(uploaded_file)
    else:
        raise ValueError(f"Unsupported file format: {os.path.splitext(filename)[1]}. Only PDF and DOCX are supported.")
